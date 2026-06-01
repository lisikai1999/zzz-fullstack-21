import io
from pathlib import Path

import pytest
from docx import Document


def _create_test_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("李明")
    doc.add_paragraph("liming@example.com")
    doc.add_paragraph("13912345678")
    doc.add_paragraph("")
    doc.add_paragraph("Education")
    doc.add_paragraph("2015-2019 浙江大学 Bachelor 计算机科学")
    doc.add_paragraph("")
    doc.add_paragraph("Work Experience")
    doc.add_paragraph("2019.06-2022.06 阿里巴巴 后端工程师")
    doc.add_paragraph("负责微服务架构设计与实现")
    doc.add_paragraph("")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, Java, Docker, Kubernetes, MySQL, Redis")

    file_path = tmp_path / "test_resume.docx"
    doc.save(file_path)
    return file_path


class TestUploadEndpoint:
    def test_upload_docx(self, client, tmp_path):
        file_path = _create_test_docx(tmp_path)
        with open(file_path, "rb") as f:
            response = client.post(
                "/api/candidates/upload",
                files={"file": ("resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        assert response.status_code == 200
        data = response.json()
        assert data["name"] == "李明"
        assert data["email"] == "liming@example.com"
        assert data["phone"] == "13912345678"
        assert len(data["skills"]) > 0

    def test_upload_invalid_format(self, client):
        response = client.post(
            "/api/candidates/upload",
            files={"file": ("test.txt", io.BytesIO(b"hello"), "text/plain")},
        )
        assert response.status_code == 400


class TestJobCRUD:
    def test_create_and_get(self, client):
        resp = client.post("/api/jobs", json={
            "title": "前端工程师",
            "department": "技术部",
            "min_education": "bachelor",
            "min_years": 2,
            "required_skills": ["JavaScript", "Vue.js"],
            "weighted_skills": [
                {"skill": "Vue.js", "weight": 8},
                {"skill": "TypeScript", "weight": 6},
            ],
        })
        assert resp.status_code == 200
        job = resp.json()
        assert job["title"] == "前端工程师"
        assert job["id"] > 0

        resp2 = client.get(f"/api/jobs/{job['id']}")
        assert resp2.status_code == 200
        assert resp2.json()["title"] == "前端工程师"

    def test_list_jobs(self, client):
        client.post("/api/jobs", json={"title": "Job A"})
        client.post("/api/jobs", json={"title": "Job B"})
        resp = client.get("/api/jobs")
        assert resp.status_code == 200
        assert len(resp.json()) >= 2

    def test_delete_job(self, client):
        resp = client.post("/api/jobs", json={"title": "To Delete"})
        job_id = resp.json()["id"]
        del_resp = client.delete(f"/api/jobs/{job_id}")
        assert del_resp.status_code == 200
        get_resp = client.get(f"/api/jobs/{job_id}")
        assert get_resp.status_code == 404


class TestMatchEndpoint:
    def test_top_candidates_for_job(self, client, tmp_path):
        file_path = _create_test_docx(tmp_path)
        with open(file_path, "rb") as f:
            client.post(
                "/api/candidates/upload",
                files={"file": ("resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )

        job_resp = client.post("/api/jobs", json={
            "title": "后端工程师",
            "min_education": "bachelor",
            "min_years": 2,
            "required_skills": ["Python"],
            "weighted_skills": [
                {"skill": "Python", "weight": 8},
                {"skill": "Docker", "weight": 6},
            ],
        })
        job_id = job_resp.json()["id"]

        match_resp = client.get(f"/api/matching/job/{job_id}/top?k=5")
        assert match_resp.status_code == 200
        results = match_resp.json()
        assert isinstance(results, list)
